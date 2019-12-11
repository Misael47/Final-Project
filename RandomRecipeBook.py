import requests
import docx

# API random recipes
first_recipe = requests.get('https://taco-1150.herokuapp.com/random/?full_taco=true').json()
second_recipe = requests.get('https://taco-1150.herokuapp.com/random/?full_taco=true').json()
third_recipe = requests.get('https://taco-1150.herokuapp.com/random/?full_taco=true').json()
#gets taco recipes from API
def getrandomrecipe(url):
    try:
        cookbook = requests.get(url).json()
        for recipe in cookbook:         # retrieves ingredients from the list API's
            base = recipe['base_layer']
            spices = recipe['seasoning']
            mix = recipe['mixin']
            additive = recipe['condiment']
            exterior = recipe['shell']
            print(f'{base}: {spices}: {mix}: {additive}: {exterior}\n')
    except:
        print('Server not found, want to taco bout it')

#this Section will retrieve 3 random taco recipes.
document = docx.Document()
# won't work without internet access
# this section will create my word document with the following items I need.
document.add_heading('Random Taco Cookbook', 0)
# Adds a heading which differs from 0-4, 0 being title size
document.add_picture('Modified_TacoPhoto.jpg')
# adds picture with the size you want to fit in your word file. Not putting any gives you original size.
document.add_heading('Credits', 1)
# headings format size smaller than 0 for 'credits'
document.add_paragraph('Photo author: Chad Montano')
# adding photo author into paragraph, along with the Url and your name as Coder
document.add_paragraph('Code: Misael Arreola')
para = document.add_paragraph('API URL: https://taco-1150.herokuapp.com/random/?full_taco=true')
#run = para.add_run()
#run.add_break(docx.text.WD_BREAK.PAGE)

document.save('Random_Taco_Cookbook.docx')